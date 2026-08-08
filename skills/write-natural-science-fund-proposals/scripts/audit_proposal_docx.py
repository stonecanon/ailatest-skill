#!/usr/bin/env python3
"""Structural audit for the user's Chinese natural-science-fund DOCX files."""

from __future__ import annotations

import argparse
from pathlib import Path
import re
import sys
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


FORBIDDEN = (
    "国家基金本子",
    "参考国家基金",
    "参考已有文件",
    "沿用已有本子",
    "根据修改指令",
    "量程与精度，需填",
    "需填",
    "待补",
    "待核实",
)
CAPTION_RE = re.compile(r"^图\s*(\d+)\s+")
NUMERIC_CITATION_RE = re.compile(r"\[(?:\d+(?:[-–—]\d+)?)(?:,\d+(?:[-–—]\d+)?)*\]")
NAME_RE = re.compile(r"翁建涛|Jiantao Weng")


def run_is_superscript(run) -> bool:
    return run.font.superscript is True


def audit(path: Path) -> tuple[list[str], list[str]]:
    errors: list[str] = []
    warnings: list[str] = []
    try:
        doc = Document(path)
    except Exception as exc:
        return [f"cannot open DOCX: {exc}"], warnings

    paragraphs = doc.paragraphs
    texts = [p.text.strip() for p in paragraphs]
    all_text = "\n".join(texts)
    for table in doc.tables:
        all_text += "\n" + "\n".join(
            cell.text for row in table.rows for cell in row.cells
        )

    for phrase in FORBIDDEN:
        count = all_text.count(phrase)
        if count:
            errors.append(f"forbidden or unresolved phrase {phrase!r}: {count}")

    captions: dict[int, object] = {}
    for paragraph in paragraphs:
        match = CAPTION_RE.match(paragraph.text.strip())
        if match:
            captions[int(match.group(1))] = paragraph
    if not captions:
        warnings.append("no numbered figure captions found")
    for number, paragraph in sorted(captions.items()):
        for run in (r for r in paragraph.runs if r.text):
            fonts = run._element.get_or_add_rPr().rFonts
            east_asia = fonts.get(qn("w:eastAsia"))
            size = run.font.size.pt if run.font.size else None
            if east_asia not in {"华文宋体", "STSong", "宋体"}:
                warnings.append(f"figure {number} caption uses Chinese font {east_asia!r}")
            if size is None or not 10.5 <= size <= 11.5:
                warnings.append(f"figure {number} caption size is {size!r}, expected about 11 pt")
            if run.bold is True:
                warnings.append(f"figure {number} caption is bold; established style is regular")

    caption_numbers = set(captions)
    reference_numbers: set[int] = set()
    for text in texts:
        if CAPTION_RE.match(text):
            continue
        for match in re.finditer(r"图\s*(\d+)", text):
            reference_numbers.add(int(match.group(1)))
    for number in sorted(caption_numbers - reference_numbers):
        warnings.append(f"figure {number} has no正文 reference outside its caption")
    for number in sorted(reference_numbers - caption_numbers):
        warnings.append(f"正文 references figure {number}, but no matching caption was found")

    reference_start = next(
        (i for i, text in enumerate(texts) if "参考文献" in text), len(paragraphs)
    )
    for paragraph in paragraphs[:reference_start]:
        for run in paragraph.runs:
            if NUMERIC_CITATION_RE.search(run.text) and not run_is_superscript(run):
                errors.append(f"numeric citation is not superscript: {run.text!r}")

    achievements_start = next(
        (i for i, text in enumerate(texts) if text.startswith("6.1.6")), None
    )
    achievements_end = next(
        (i for i, text in enumerate(texts) if text.startswith("6.1.7") or text.startswith("6.2")),
        None,
    )
    if achievements_start is not None and achievements_end is not None:
        for paragraph in paragraphs[achievements_start + 1 : achievements_end]:
            for run in paragraph.runs:
                if NAME_RE.search(run.text) and run.bold is not True:
                    errors.append(f"applicant name is not bold in achievements: {run.text!r}")

    heading_index = next(
        (i for i, text in enumerate(texts) if text.startswith("6.1.7")), None
    )
    if heading_index is not None:
        heading = paragraphs[heading_index]
        expected = "6.1.7 知识产权与设备方法基础。"
        if heading.text.strip() != expected:
            errors.append("6.1.7 heading is not on its own exact line")
        if not all(run.bold is True for run in heading.runs if run.text):
            errors.append("6.1.7 heading is not fully bold")
        if heading_index + 1 >= len(paragraphs) or not paragraphs[heading_index + 1].text.strip().startswith(
            "围绕建筑环境感知、人员行为记录、建筑性能评价和运行调控"
        ):
            errors.append("6.1.7 explanatory paragraph is missing or not on the next line")

    sensor_cluster = all(term in all_text for term in ("温湿度", "CO2", "PM2.5", "照度"))
    if sensor_cluster and "噪声" not in all_text:
        errors.append("environmental sensor list omits noise level")

    software_lines = [text for text in texts if "软件著作权" in text and "申请" in text]
    for text in software_lines:
        if "2项" not in text:
            warnings.append(f"software-copyright count may be inconsistent: {text}")

    if len(doc.inline_shapes) != len(captions):
        warnings.append(
            f"inline figure count {len(doc.inline_shapes)} differs from caption count {len(captions)}"
        )

    try:
        with ZipFile(path) as archive:
            bad = archive.testzip()
            names = set(archive.namelist())
            xml = archive.read("word/document.xml")
            if bad:
                errors.append(f"corrupt ZIP member: {bad}")
            if xml.count(b"<w:ins") or xml.count(b"<w:del"):
                errors.append("tracked insertions or deletions remain")
            if "word/comments.xml" in names:
                warnings.append("reviewer comments remain")
    except Exception as exc:
        errors.append(f"DOCX package audit failed: {exc}")

    return errors, warnings


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()
    errors, warnings = audit(args.docx)
    for warning in warnings:
        print(f"warning: {warning}")
    for error in errors:
        print(f"error: {error}")
    print(
        f"{'FAIL' if errors else 'PASS'}: {args.docx} "
        f"({len(errors)} errors, {len(warnings)} warnings)"
    )
    return 1 if errors else 0


if __name__ == "__main__":
    sys.exit(main())
